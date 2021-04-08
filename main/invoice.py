from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx2pdf import convert


def create_invoice(d):
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)

    paragraph1 = document.add_paragraph()
    run1 = paragraph1.add_run()
    run1.add_picture('main/Invoices/logo1.png')
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph2 = document.add_paragraph()
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = paragraph2.add_run()
    run2.add_text('\nGENESIS TRAINING - INVOICE')
    run2.font.size = Pt(18)

    s = '\nName : {0}\nBank Account No. : {1}\nIFSC : {2}\n'.format(d['trainer_name'], d['acc'], d['ifsc'])
    s += 'PAN : {0}\nBank Name : {1}\nPhone Number : {2}\n'.format(d['pan'], d['bank'], d['phno'])
    s += 'Email : {0}\nBase Location : {1}\n'.format(d['email'], d['base_location'])
    paragraph3 = document.add_paragraph()
    paragraph3.paragraph_format.line_spacing = 1.5
    run3 = paragraph3.add_run(s)
    run3.font.size = Pt(13)

    no_days = int(d['no_of_days']) + 2
    fees_per_day = d['pay_per_day']
    travel_allowance = d['travel']
    food_allowance = d['food']

    table = document.add_table(rows=no_days, cols=5)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Date'
    table.cell(0, 1).text = 'College'
    table.cell(0, 2).text = 'Fees/day'
    table.cell(0, 3).text = 'Travel Allowance'
    table.cell(0, 4).text = 'Food Allowance'
    table.cell(no_days - 1, 1).text = 'Total'
    table.cell(no_days - 1, 2).text = str(fees_per_day * (no_days-2))
    total = fees_per_day * (no_days-2)

    for i in range(1, no_days - 1):
        table.cell(i, 3).text = 'NA'
    table.cell(no_days - 1, 3).text = '0'

    if travel_allowance == 'Yes':
        if no_days > 3:
            table.cell(1, 3).text = '500'
            table.cell(no_days-2, 3).text = '500'
            table.cell(no_days-1, 3).text = '1000'
        if no_days == 3:
            table.cell(1, 3).text = '1000'
        table.cell(no_days - 1, 3).text = '1000'
        total += 1000


    if food_allowance != 'No':
        table.cell(no_days - 1, 4).text = str((no_days-2) * 150)
        total += (no_days-2) * 150
    else:
        table.cell(no_days - 1, 4).text = str(0)

    for i in range(1, no_days - 1):
        table.cell(i,0).text = d['dates'][i-1]
        table.cell(i, 1).text = d['college_name']
        table.cell(i, 2).text = str(fees_per_day)
        if food_allowance != 'No':
            table.cell(i, 4).text = str(150)
        else:
            table.cell(i, 4).text = 'NA'

    def make_rows_bold(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                     for run in paragraph.runs:
                        run.font.bold = True

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                 for run in paragraph.runs:
                    run.font.size = Pt(13)

    make_rows_bold(table.rows[0], table.rows[-1])

    paragraph4 = document.add_paragraph()
    paragraph4.paragraph_format.line_spacing = 1.75
    run4 = paragraph4.add_run('\nTotal: {0}'.format(total))
    run4.font.size = Pt(13)
    run4.bold = True

    path = 'main/Invoices/Invoice.docx'
    document.save(path)
    convert(path)
    return path.split('.')[0]+'.pdf'
